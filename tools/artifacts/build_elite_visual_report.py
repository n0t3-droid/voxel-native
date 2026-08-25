#!/usr/bin/env python3
"""Build a DOCX dossier from one explicit canonical QA evidence manifest."""

from __future__ import annotations

import argparse
import json
import os
import uuid
from pathlib import Path
from typing import Any, Sequence

from evidence_manifest_consumer import (
    CanonicalEvidence,
    EvidenceContractError,
    iter_claims,
    iter_issues,
    load_canonical_evidence,
    publish_no_clobber,
    validate_output_path,
    validation_summary,
    verified_screenshots,
)


NAVY = "17365D"
BLUE = "2F75B5"
TEAL = "2D8C85"
LIGHT_BLUE = "DCE6F1"
LIGHT_GREEN = "E2F0D9"
LIGHT_GRAY = "F2F2F2"
DARK = "1F2937"
MUTED = "5F6B78"
WHITE = "FFFFFF"
PAGE_WIDTH_DXA = 9360
MAX_ARTIFACT_ROWS = 300
FILE_IDENTITY_HEADERS = ("Kind", "Path", "Bytes", "SHA-256")
FILE_IDENTITY_WIDTHS_DXA = (1000, 3860, 1100, 3400)
FILE_IDENTITY_ROWS_PER_PAGE = 8
COHORT_KIND_NAMES = (
    "NaturalGrove",
    "NaturalKarst",
    "NaturalMesa",
    "AstralCrystal",
    "AstralBasalt",
    "AstralReef",
)


def load_docx_dependencies() -> None:
    """Load bundled document dependencies only for a real artifact build."""

    global Document, WD_ORIENT, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    global WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor
    try:
        from docx import Document as _Document
        from docx.enum.section import WD_ORIENT as _WD_ORIENT
        from docx.enum.table import (
            WD_CELL_VERTICAL_ALIGNMENT as _WD_CELL_VERTICAL_ALIGNMENT,
            WD_TABLE_ALIGNMENT as _WD_TABLE_ALIGNMENT,
        )
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement as _OxmlElement
        from docx.oxml.ns import qn as _qn
        from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGBColor
    except ImportError as error:
        raise EvidenceContractError(
            "python-docx is unavailable; run a real build with the bundled workspace dependency runtime"
        ) from error
    Document = _Document
    WD_ORIENT = _WD_ORIENT
    WD_CELL_VERTICAL_ALIGNMENT = _WD_CELL_VERTICAL_ALIGNMENT
    WD_TABLE_ALIGNMENT = _WD_TABLE_ALIGNMENT
    WD_ALIGN_PARAGRAPH = _WD_ALIGN_PARAGRAPH
    OxmlElement = _OxmlElement
    qn = _qn
    Inches = _Inches
    Pt = _Pt
    RGBColor = _RGBColor


def set_cell_fill(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != PAGE_WIDTH_DXA:
        raise EvidenceContractError("DOCX table widths must total exactly 9360 DXA")
    table.autofit = False
    properties = table._tbl.tblPr
    table_width = properties.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")
    table_indent = properties.find(qn("w:tblInd"))
    if table_indent is None:
        table_indent = OxmlElement("w:tblInd")
        properties.append(table_indent)
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(width))
            cell_width.set(qn("w:type"), "dxa")


def keep_row_together(row, *, repeat_header: bool = False) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))
    if repeat_header and properties.find(qn("w:tblHeader")) is None:
        properties.append(OxmlElement("w:tblHeader"))


def keep_table_as_block(table) -> None:
    """Link every table row except the last through Word keep-next paragraphs."""

    for row in table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph, generated_date: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(f"VOXEL-NATIVE  |  CANONICAL QA EVIDENCE  |  {generated_date}  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_document(document: Document, generated_date: str) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 11, 9, 4),
    ):
        style = document.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE if name == "Heading 1" else NAVY)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left, right = header_table.rows[0].cells
    left.text = "VOXEL-NATIVE"
    right.text = "EVIDENCE DOSSIER"
    right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in (left, right):
        set_cell_margins(cell, top=0, bottom=40)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(NAVY)
    set_table_geometry(header_table, [4680, 4680])
    add_page_number(section.footer.paragraphs[0], generated_date)


def add_kicker(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text.upper())
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(TEAL)


def add_callout(document: Document, label: str, text: str, fill: str = LIGHT_GREEN) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.keep_together = True
    label_run = paragraph.add_run(label + "  ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(NAVY)
    paragraph.add_run(text)
    keep_row_together(table.rows[0])
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(document: Document, items: Sequence[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.28)
        paragraph.paragraph_format.first_line_indent = Inches(-0.18)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run(item)


def add_matrix(
    document: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[object]],
    widths_dxa: Sequence[int],
    *,
    keep_as_block: bool = False,
) -> Any:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(WHITE)
    keep_row_together(table.rows[0], repeat_header=True)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        keep_row_together(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                set_cell_fill(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(str(value))
            run.font.size = Pt(8)
    if keep_as_block:
        keep_table_as_block(table)
    set_table_geometry(table, widths_dxa)
    document.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def fmt_number(value: object, digits: int = 2) -> str:
    if type(value) is int:
        return f"{value:,}"
    if type(value) is float:
        return f"{value:,.{digits}f}"
    return str(value)


def display_hash(value: str) -> str:
    return " ".join(value[index : index + 8] for index in range(0, len(value), 8))


def fmt_optional_count(value: object) -> str:
    """Preserve an upstream null instead of turning missing work into zero."""

    return "not recorded" if value is None else fmt_number(value, 0)


def route_evidence_text(route: dict[str, Any]) -> str:
    return (
        f"{route['requested_route_focus']} -> {route['resolved_route_focus']}\n"
        f"{fmt_number(route['requested_route_distance_m'], 0)} m"
    )


def generation_identity_text(identity: dict[str, Any]) -> str:
    return (
        f"{identity.get('world_name') or 'unrecorded'}\n"
        f"seed {identity.get('world_seed')}; profile {identity.get('world_profile')}; "
        f"scenery {identity.get('scenery_quality')}; grammar {identity['terrain_grammar']}"
    )


def edit_store_state_text(edit_store: dict[str, Any]) -> str:
    return (
        f"{edit_store['world_edit_store_status']}; compatible "
        f"{'yes' if edit_store['world_edit_store_compatible'] else 'no'}\n"
        f"edited chunks {fmt_optional_count(edit_store['world_edit_store_edited_chunks'])}; "
        f"reason {edit_store['world_edit_store_block_reason_code'] or 'none'}"
    )


def edit_store_identity_text(edit_store: dict[str, Any]) -> str:
    return (
        f"seed {edit_store['world_edit_store_seed']}; "
        f"profile {edit_store['world_edit_store_profile']}; "
        f"scenery {edit_store['world_edit_store_scenery_quality']}; "
        f"grammar {edit_store['world_edit_store_terrain_grammar']}"
    )


def cohort_kind_text(counts: Sequence[int]) -> str:
    return "; ".join(
        f"{name} {fmt_number(count, 0)}"
        for name, count in zip(COHORT_KIND_NAMES, counts, strict=True)
    )


def plan_balanced_table_pages(
    row_count: int, maximum_per_page: int = FILE_IDENTITY_ROWS_PER_PAGE
) -> tuple[int, ...]:
    """Return balanced page populations without a one-row multipage orphan."""

    if type(row_count) is not int or row_count < 1:
        raise ValueError("table pagination requires a positive integer row count")
    if type(maximum_per_page) is not int or maximum_per_page < 3:
        raise ValueError("table pagination requires a per-page maximum of at least three")
    page_count = (row_count + maximum_per_page - 1) // maximum_per_page
    base, remainder = divmod(row_count, page_count)
    page_sizes = tuple(base + (1 if index < remainder else 0) for index in range(page_count))
    if page_count > 1 and min(page_sizes) < 2:
        raise EvidenceContractError("table pagination would create a one-row orphan page")
    return page_sizes


def balanced_row_chunks(
    rows: Sequence[Sequence[object]],
    maximum_per_page: int = FILE_IDENTITY_ROWS_PER_PAGE,
) -> list[list[Sequence[object]]]:
    """Partition rows exactly once, in order, according to the balanced plan."""

    page_sizes = plan_balanced_table_pages(len(rows), maximum_per_page)
    chunks: list[list[Sequence[object]]] = []
    start = 0
    for page_size in page_sizes:
        chunks.append(list(rows[start : start + page_size]))
        start += page_size
    if start != len(rows):
        raise EvidenceContractError("table pagination did not consume every row exactly once")
    return chunks


def file_identity_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    return [
        [
            record["kind"],
            record["path"],
            fmt_number(record["size_bytes"], 0),
            display_hash(record["sha256"]),
        ]
        for record in evidence.data["file_hashes"]
    ]


def add_file_identity_pages(document: Document, evidence: CanonicalEvidence) -> None:
    """Emit balanced identity tables at deterministic page boundaries."""

    chunks = balanced_row_chunks(file_identity_rows(evidence))
    for index, chunk in enumerate(chunks):
        heading_text = "Evidence file identity"
        if len(chunks) > 1:
            heading_text += f" ({index + 1}/{len(chunks)})"
        heading = document.add_heading(heading_text, level=1)
        heading.paragraph_format.page_break_before = True
        add_matrix(
            document,
            FILE_IDENTITY_HEADERS,
            chunk,
            FILE_IDENTITY_WIDTHS_DXA,
            keep_as_block=True,
        )


def run_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    rows: list[list[str]] = []
    for run in evidence.runs:
        observations = run["raw_observations"]
        identity = observations["run_identity"]
        viewport = observations["viewport"]
        route = observations["route"]
        frame = observations["route_frame_times"]
        planetary = observations["planetary_streaming"]
        telemetry = planetary["telemetry"]
        rows.append(
            [
                run["input_path"],
                f"{identity.get('build_profile')} / {identity.get('world_profile', 'unrecorded')} / {identity['terrain_grammar']}",
                f"{viewport['physical_width']}x{viewport['physical_height']} @ {fmt_number(viewport['dpi_percent'], 0)}%",
                route_evidence_text(route),
                f"n={fmt_number(frame['sample_count'], 0)}; p50 {fmt_number(frame['median_ms'])}; p95 {fmt_number(frame['p95_ms'])}; p99 {fmt_number(frame['p99_ms'])}; max {fmt_number(frame['max_ms'])} ms",
                f"{telemetry['surface_material_mode']} / {planetary['live']['profile']}; Hydro {telemetry['hydro_mode']}; cohorts {telemetry['semantic_cohort_mode']}",
                str(len(observations["screenshots"]["referenced_files"])),
            ]
        )
    return rows


def generation_identity_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    rows: list[list[str]] = []
    for run in evidence.runs:
        observations = run["raw_observations"]
        identity = observations["run_identity"]
        edit_store = observations["world_edit_store"]
        telemetry = observations["planetary_streaming"]["telemetry"]
        rows.append(
            [
                run["input_path"],
                generation_identity_text(identity),
                edit_store_state_text(edit_store),
                edit_store_identity_text(edit_store),
                f"desired {telemetry['desired_terrain_grammar']} -> active {telemetry['active_terrain_grammar']}",
            ]
        )
    return rows


def route_detail_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    rows: list[list[str]] = []
    for run in evidence.runs:
        route = run["raw_observations"]["route"]
        rows.append(
            [
                run["input_path"],
                f"{route['requested_route_focus']} -> {route['resolved_route_focus']}",
                f"available {'yes' if route['route_focus_available'] else 'no'}; reason {route['route_focus_unavailable_reason'] or 'none'}",
                "not recorded" if route["route_focus_anchor"] is None else str(route["route_focus_anchor"]),
                f"candidates {fmt_optional_count(route['route_focus_search_visited_candidates'])}/"
                f"{fmt_number(route['route_focus_search_candidate_cap'], 0)}; classifications "
                f"{fmt_optional_count(route['route_focus_classification_queries'])}/"
                f"{fmt_number(route['route_focus_classification_query_cap'], 0)}; cap exhausted "
                f"{'yes' if route['route_focus_search_cap_exhausted'] else 'no'}",
            ]
        )
    return rows


def layer_evidence_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    """Render per-kind Hydro and semantic-cohort truth without inferring quality."""

    rows: list[list[str]] = []
    for run in evidence.runs:
        planetary = run["raw_observations"]["planetary_streaming"]
        live = planetary["live"]
        budgets = planetary["budgets"]
        telemetry = planetary["telemetry"]
        rows.append(
            [
                run["input_path"],
                "Far Hydro",
                telemetry["hydro_mode"],
                f"water {fmt_number(live['resident_water_indices'], 0)} indices, rings {live['water_ring_indices']}; "
                f"lava {fmt_number(live['resident_lava_indices'], 0)} indices, rings {live['lava_ring_indices']}",
                f"entities {fmt_number(live['resident_fluid_entities'], 0)}; vertices {fmt_number(live['resident_fluid_vertices'], 0)}; "
                f"indices {fmt_number(live['resident_fluid_indices'], 0)}; bytes {fmt_number(live['resident_fluid_mesh_bytes'], 0)}",
                f"mesh {fmt_number(budgets['budget_fluid_mesh_bytes'], 0)} B; Hydro atomic {fmt_number(budgets['budget_hydro_atomic_ring_build_bytes'], 0)} B; "
                f"observation/integrity {'Passed' if telemetry['resident_fluid_observation_valid'] and telemetry['resident_fluid_kind_integrity_valid'] else 'Rejected'}",
            ]
        )
        rows.append(
            [
                run["input_path"],
                "Semantic cohorts",
                telemetry["semantic_cohort_mode"],
                cohort_kind_text(live["resident_semantic_cohort_kind_counts"]),
                f"cohorts {fmt_number(live['resident_semantic_cohort_count'], 0)}; entities {fmt_number(live['resident_semantic_cohort_entities'], 0)}; "
                f"vertices {fmt_number(live['resident_semantic_cohort_vertices'], 0)}; indices {fmt_number(live['resident_semantic_cohort_indices'], 0)}; "
                f"bytes {fmt_number(live['resident_semantic_cohort_mesh_bytes'], 0)}",
                f"mesh {fmt_number(budgets['budget_semantic_cohort_mesh_bytes'], 0)} B; candidates {fmt_number(telemetry['last_semantic_cohort_candidates'], 0)}/"
                f"{fmt_number(budgets['budget_semantic_cohort_hash_scans'], 0)} scans; observation/integrity "
                f"{'Passed' if telemetry['resident_semantic_cohort_observation_valid'] and telemetry['resident_semantic_cohort_payload_integrity_valid'] else 'Rejected'}",
            ]
        )
    return rows


def budget_rows(evidence: CanonicalEvidence) -> list[list[str]]:
    pairs = (
        ("Entities", "resident_entities", "budget_entities"),
        ("Vertices", "resident_vertices", "budget_vertices"),
        ("Indices", "resident_indices", "budget_indices"),
        ("Mesh bytes", "resident_mesh_bytes", "budget_mesh_bytes"),
        ("Sample-cache bytes", "live_sample_cache_bytes", "budget_sample_cache_bytes"),
    )
    rows: list[list[str]] = []
    for run in evidence.runs:
        planetary = run["raw_observations"]["planetary_streaming"]
        live = planetary["live"]
        budgets = planetary["budgets"]
        telemetry = planetary["telemetry"]
        for label, live_field, budget_field in pairs:
            current = int(live[live_field])
            budget = int(budgets[budget_field])
            usage = 0.0 if budget == 0 and current == 0 else current / budget if budget else float("inf")
            rows.append(
                [
                    run["input_path"],
                    label,
                    fmt_number(current, 0),
                    fmt_number(budget, 0),
                    f"{usage:.1%}" if usage != float("inf") else "undefined",
                    "Passed",
                ]
            )
        rows.append(
            [
                run["input_path"],
                "Peak sample-cache bytes",
                fmt_number(telemetry["peak_live_sample_cache_bytes"], 0),
                fmt_number(budgets["budget_sample_cache_bytes"], 0),
                f"{telemetry['peak_live_sample_cache_bytes'] / budgets['budget_sample_cache_bytes']:.1%}"
                if budgets["budget_sample_cache_bytes"]
                else "undefined",
                "Passed",
            ]
        )
    return rows


def build_report(
    evidence: CanonicalEvidence,
    output_path: Path,
    repo_root: Path,
) -> Path:
    load_docx_dependencies()
    claims = list(iter_claims(evidence))
    issues = list(iter_issues(evidence))
    if len(claims) > MAX_ARTIFACT_ROWS or len(issues) > MAX_ARTIFACT_ROWS:
        raise EvidenceContractError("claim or issue ledger exceeds the DOCX artifact row cap")
    if len(evidence.data["file_hashes"]) > MAX_ARTIFACT_ROWS:
        raise EvidenceContractError("file hash ledger exceeds the DOCX artifact row cap")
    screenshots = verified_screenshots(evidence, repo_root)

    generated_date = evidence.generated_at.date().isoformat()
    document = Document()
    configure_document(document, generated_date)
    properties = document.core_properties
    properties.title = "Voxel-Native Canonical QA Evidence Dossier"
    properties.author = "Voxel-Native evidence pipeline"
    properties.subject = "Manifest-backed visual, route, and planetary-streaming evidence"
    properties.keywords = "Voxel-Native, QA, evidence manifest, planetary streaming"
    properties.created = evidence.generated_at.replace(tzinfo=None)
    properties.modified = evidence.generated_at.replace(tzinfo=None)

    add_kicker(document, "Canonical evidence / no inferred results")
    title = document.add_paragraph(style="Title")
    title.add_run("Voxel-Native QA\nevidence dossier")
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run(
        "A bounded rendering of one explicit evidence manifest - not a scan of latest runs."
    )
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor.from_string(MUTED)

    add_matrix(
        document,
        ["GENERATED", "CLASSIFICATION", "RUNS", "MANIFEST SHA-256"],
        [[generated_date, evidence.data["overall_classification"], len(evidence.runs), evidence.manifest_sha256]],
        [1500, 1800, 900, 5160],
    )
    if screenshots:
        run, display, image_path, record = screenshots[0]
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run().add_picture(str(image_path), width=Inches(6.45))
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(9)
        caption_run = caption.add_run(
            f"Manifest-referenced PNG: {display} | run {run['input_path']} | sha256 {display_hash(record['sha256'])}"
        )
        caption_run.italic = True
        caption_run.font.size = Pt(7.5)
        caption_run.font.color.rgb = RGBColor.from_string(MUTED)

    document.add_heading("Evidence boundary", level=1)
    document.add_paragraph(
        "Every run below uses QA report schema 2.6.0 through manifest schema 1.6.0, including immutable terrain grammar, combined dense-residency proof, "
        "compatible edit-store identity, route-only frame-time quantiles, explicit viewport provenance, planetary live "
        "values and hard budgets, and manifest-referenced PNG identities. The aggregate classification is "
        "Observed because measured runtime values are observations even when integrity and budget checks Passed."
    )
    add_callout(
        document,
        "No fabricated release result",
        "This manifest contains no automated test-suite transcript or test total. This dossier therefore reports none. "
        "PNG completion and hashes prove byte identity and container completion, not perceptual visual quality.",
        LIGHT_BLUE,
    )

    document.add_heading("Run evidence", level=1)
    add_matrix(
        document,
        ["Explicit run", "Build / world / grammar", "Viewport", "Route", "Route-only frame time", "Far surface", "PNGs"],
        run_rows(evidence),
        [1650, 1250, 1100, 1300, 2100, 1500, 460],
    )

    document.add_heading("Generation identity and edit-store compatibility", level=1)
    document.add_paragraph(
        "Terrain grammar is part of the immutable world identity. The edit store repeats that exact identity, while far-field desired and active grammar remain separate serialized observations. Compatible is a manifest-validated authority state, not an inference from an empty directory or world name."
    )
    add_matrix(
        document,
        ["Run", "World generation identity", "Edit-store state", "Edit-store identity", "Far grammar"],
        generation_identity_rows(evidence),
        [1800, 2200, 1700, 2200, 1460],
    )

    document.add_heading("Route resolution and bounded search work", level=1)
    document.add_paragraph(
        "Requested and resolved focus remain separate. Actual search counters are rendered as not recorded when the upstream API serialized null; null is never rewritten as zero."
    )
    add_matrix(
        document,
        ["Run", "Requested -> resolved", "Availability / reason", "Anchor", "Actual / cap"],
        route_detail_rows(evidence),
        [1700, 1550, 1650, 1200, 3260],
    )

    document.add_heading("Far Hydro and semantic-cohort evidence", level=1)
    document.add_paragraph(
        "Counts below are serialized post-deferred observations from the explicit runs. Water and lava remain separate, and the six semantic kinds retain their fixed schema order. Passed means only that the manifest's bounded observation, scheduler, payload, and budget invariants validated; it is not visual acceptance."
    )
    add_matrix(
        document,
        ["Run", "Layer", "Mode", "Kinds / rings", "Observed payload", "Budget / integrity"],
        layer_evidence_rows(evidence),
        [1450, 1000, 1050, 2250, 1900, 1710],
    )

    document.add_heading("Serialized hard-budget evidence", level=1)
    document.add_paragraph(
        "These comparisons repeat the manifest's like-for-like planetary checks. They do not imply visual acceptance or a causal performance gain."
    )
    add_matrix(
        document,
        ["Run", "Measure", "Live / peak", "Budget", "Usage", "Manifest decision"],
        budget_rows(evidence),
        [2000, 1850, 1400, 1400, 1000, 1710],
    )

    document.add_heading("Claim ledger", level=1)
    add_matrix(
        document,
        ["Scope", "Classification", "Claim", "Evidence paths"],
        [
            [scope, claim["classification"], claim["statement"], "\n".join(claim["evidence"]) or "None"]
            for scope, claim in claims
        ],
        [1650, 1100, 3450, 3160],
    )

    document.add_heading("Issue ledger", level=1)
    if issues:
        add_matrix(
            document,
            ["Scope", "Classification", "Code / field", "Recorded message"],
            [
                [scope, issue["classification"], f"{issue['code']} / {issue['field']}", issue["message"]]
                for scope, issue in issues
            ],
            [1650, 1100, 2500, 4110],
        )
    else:
        document.add_paragraph("The manifest records no issues for this evidence set.")

    add_file_identity_pages(document, evidence)

    document.add_heading("Standing interpretation limits", level=1)
    add_bullets(
        document,
        [
            "One manifest run records one viewport. It does not imply completion of the responsive viewport and DPI matrix.",
            "Average FPS and route quantiles describe the serialized route on the recorded build and hardware; they are not an inferred universal threshold or an A/B uplift.",
            "A PNG hash proves exact bytes. Human visual inspection must still check clipping, overlap, holes, repetition, lighting, transitions, and motion.",
            "File hashes do not prove authorship, an unrecorded Git revision, or correspondence with source outside the serialized provenance fields.",
            "Automated test counts require a separately hashed gate transcript and are deliberately absent here.",
        ],
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    temporary = output_path.with_name(f".{output_path.name}.{os.getpid()}.{uuid.uuid4().hex}.partial")
    document.save(temporary)
    publish_no_clobber(temporary, output_path)
    return output_path


def parse_args(argv: Sequence[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--evidence-manifest", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=Path(__file__).resolve().parents[2],
        help="repository root used only to resolve manifest-relative evidence paths",
    )
    parser.add_argument(
        "--check-only",
        action="store_true",
        help="validate the manifest, output path, and referenced PNG bytes without writing",
    )
    return parser.parse_args(argv)


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv)
    try:
        repo_root = args.repo_root.resolve(strict=False)
        evidence = load_canonical_evidence(args.evidence_manifest)
        output = validate_output_path(args.output, repo_root, ".docx")
        if args.check_only:
            verified_screenshots(evidence, repo_root)
        else:
            build_report(evidence, output, repo_root)
    except (EvidenceContractError, OSError, ValueError) as error:
        print(f"DOCX artifact rejected: {error}", file=os.sys.stderr)
        return 2
    print(json.dumps(validation_summary(evidence, output), sort_keys=True))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
